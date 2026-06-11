"""Markdown -> formatted .docx converter for this project's documents.

Replaces the content-hardcoded build_paper.py approach: the markdown file is
the single source of truth; this script renders it with academic formatting
(cover page, justified Times New Roman body, navy Calibri headings, styled
tables, embedded figures, code blocks, page-number footer).

Usage (from project root):
    python docs/build_docx.py docs/RESEARCH_PAPER.md
    python docs/build_docx.py docs/THESIS.md
    python docs/build_docx.py            # builds both if present

Supported markdown subset: #..#### headings, paragraphs, - bullets,
1. numbered lists, > quotes, ``` code fences, | tables |, ![img](path),
**bold** / *italic* / `code` inline, --- rules.
"""
from __future__ import annotations

import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
DARK_NAVY = RGBColor(0x0D, 0x1F, 0x3D)
INK = RGBColor(0x10, 0x10, 0x10)
GREY = RGBColor(0x55, 0x5B, 0x63)

BODY_FONT = "Times New Roman"
HEAD_FONT = "Calibri"
CODE_FONT = "Consolas"


# ---------------------------------------------------------------------------
# Low-level docx helpers
# ---------------------------------------------------------------------------

def _page_number_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Page ")
    run.font.name = HEAD_FONT
    run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def _table_borders(table) -> None:
    tbl = table._tbl
    pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "9AA4B2")
        borders.append(el)
    pr.append(borders)


def _shade_cell(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


INLINE_RE = re.compile(
    r"(\*\*\*(?P<bi>.+?)\*\*\*)|(\*\*(?P<b>.+?)\*\*)|(\*(?P<i>[^*]+?)\*)"
    r"|(`(?P<c>[^`]+?)`)|(\[(?P<lt>[^\]]+)\]\((?P<lu>[^)]+)\))"
    r"|(<(?P<au>https?://[^>]+)>)"
)


def _add_inline(par, text: str, size=Pt(11), color=INK, base_font=BODY_FONT) -> None:
    """Render **bold**, *italic*, `code`, [text](url), <url> inside a paragraph."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            r = par.add_run(text[pos:m.start()])
            r.font.name = base_font; r.font.size = size; r.font.color.rgb = color
        if m.group("bi") is not None:
            r = par.add_run(m.group("bi")); r.bold = True; r.italic = True
        elif m.group("b") is not None:
            r = par.add_run(m.group("b")); r.bold = True
        elif m.group("i") is not None:
            r = par.add_run(m.group("i")); r.italic = True
        elif m.group("c") is not None:
            r = par.add_run(m.group("c")); r.font.name = CODE_FONT
        elif m.group("lt") is not None:
            r = par.add_run(m.group("lt"))
            r.font.color.rgb = RGBColor(0x05, 0x63, 0xC1); r.underline = True
        elif m.group("au") is not None:
            r = par.add_run(m.group("au"))
            r.font.color.rgb = RGBColor(0x05, 0x63, 0xC1); r.underline = True
        else:  # pragma: no cover
            r = par.add_run(m.group(0))
        r.font.size = size
        if r.font.name is None or r.font.name == BODY_FONT:
            r.font.name = base_font
        if r.font.color.rgb is None:
            r.font.color.rgb = color
        pos = m.end()
    if pos < len(text):
        r = par.add_run(text[pos:])
        r.font.name = base_font; r.font.size = size; r.font.color.rgb = color


def _body_par(doc, text, justify=True, size=Pt(11), space_after=6, indent=None):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    fmt.line_spacing = 1.5
    fmt.space_after = Pt(space_after)
    if indent:
        fmt.left_indent = indent
    _add_inline(p, text, size=size)
    return p


def _heading(doc, text, level):
    sizes = {1: 16, 2: 13, 3: 11.5, 4: 11}
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(14 if level <= 2 else 10)
    fmt.space_after = Pt(6)
    fmt.keep_with_next = True
    r = p.add_run(text)
    r.font.name = HEAD_FONT
    r.font.size = Pt(sizes.get(level, 11))
    r.bold = True
    r.italic = level >= 3
    r.font.color.rgb = NAVY if level > 1 else DARK_NAVY
    return p


def _code_block(doc, lines):
    for ln in lines:
        p = doc.add_paragraph()
        fmt = p.paragraph_format
        fmt.left_indent = Inches(0.3)
        fmt.space_after = Pt(0)
        fmt.line_spacing = 1.0
        r = p.add_run(ln if ln else " ")
        r.font.name = CODE_FONT
        r.font.size = Pt(9.5)
        r.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _table(doc, header_cells, body_rows):
    table = doc.add_table(rows=1 + len(body_rows), cols=len(header_cells))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_borders(table)
    for j, txt in enumerate(header_cells):
        cell = table.rows[0].cells[j]
        _shade_cell(cell, "1F3A5F")
        p = cell.paragraphs[0]
        _add_inline(p, txt.strip(), size=Pt(9.5),
                    color=RGBColor(0xFF, 0xFF, 0xFF), base_font=HEAD_FONT)
        for r in p.runs:
            r.bold = True
    for i, row in enumerate(body_rows, start=1):
        for j in range(len(header_cells)):
            txt = row[j].strip() if j < len(row) else ""
            p = table.rows[i].cells[j].paragraphs[0]
            _add_inline(p, txt, size=Pt(9.5))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _image(doc, path, alt):
    full = path if os.path.isabs(path) else os.path.join(HERE, path.replace("figures/", "figures" + os.sep))
    if not os.path.exists(full):
        _body_par(doc, f"[figure not found: {path}]", justify=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(full, width=Inches(6.2))
    if alt:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(alt)
        r.font.name = HEAD_FONT; r.font.size = Pt(9); r.italic = True
        r.font.color.rgb = GREY


# ---------------------------------------------------------------------------
# Document assembly
# ---------------------------------------------------------------------------

def _cover_page(doc, title, meta_lines):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.name = HEAD_FONT; r.font.size = Pt(22); r.bold = True
    r.font.color.rgb = DARK_NAVY
    doc.add_paragraph()
    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_inline(p, line, size=Pt(12), base_font=HEAD_FONT)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def convert(md_path: str, out_path: str | None = None) -> str:
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    section = doc.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(1))
    _page_number_footer(doc)

    # Title = first H1; meta = subsequent "**Key:** value" lines before first ---
    title = os.path.basename(md_path)
    i = 0
    while i < len(lines) and not lines[i].startswith("# "):
        i += 1
    meta_lines = []
    if i < len(lines):
        title = lines[i][2:].strip()
        i += 1
        while i < len(lines) and not lines[i].startswith("---") and not lines[i].startswith("## "):
            if lines[i].strip():
                meta_lines.append(lines[i].strip())
            i += 1
    _cover_page(doc, title, meta_lines)

    in_code = False
    code_buf: list[str] = []
    table_buf: list[str] = []

    def flush_table():
        nonlocal table_buf
        if len(table_buf) >= 2:
            rows = [[c for c in r.strip().strip("|").split("|")] for r in table_buf]
            header = rows[0]
            body = [r for r in rows[2:]] if len(rows) > 2 else []
            _table(doc, header, body)
        table_buf = []

    while i < len(lines):
        ln = lines[i]
        i += 1

        if in_code:
            if ln.strip().startswith("```"):
                _code_block(doc, code_buf)
                code_buf, in_code = [], False
            else:
                code_buf.append(ln)
            continue
        if ln.strip().startswith("```"):
            in_code = True
            continue

        if ln.strip().startswith("|"):
            table_buf.append(ln)
            continue
        elif table_buf:
            flush_table()

        s = ln.strip()
        if not s:
            continue
        if s == "---":
            continue

        img = re.match(r"^!\[(?P<alt>[^\]]*)\]\((?P<src>[^)]+)\)$", s)
        if img:
            _image(doc, img.group("src"), img.group("alt"))
            continue

        if s.startswith("#### "):
            _heading(doc, s[5:], 4)
        elif s.startswith("### "):
            _heading(doc, s[4:], 3)
        elif s.startswith("## "):
            _heading(doc, s[3:], 2)
        elif s.startswith("# "):
            _heading(doc, s[2:], 1)
        elif s.startswith("> "):
            p = _body_par(doc, s[2:], indent=Inches(0.3))
            for r in p.runs:
                r.italic = True
        elif re.match(r"^[-*] ", s):
            _body_par(doc, "•  " + s[2:], indent=Inches(0.25), space_after=3)
        elif re.match(r"^\d+\. ", s):
            _body_par(doc, s, indent=Inches(0.25), space_after=3)
        else:
            # join soft-wrapped continuation lines into one paragraph
            buf = [s]
            while i < len(lines):
                nxt = lines[i].strip()
                if (not nxt or nxt.startswith(("#", "|", "```", "- ", "* ", "> ", "!["))
                        or re.match(r"^\d+\. ", nxt) or nxt == "---"):
                    break
                buf.append(nxt)
                i += 1
            _body_par(doc, " ".join(buf))

    if table_buf:
        flush_table()
    if in_code and code_buf:
        _code_block(doc, code_buf)

    out = out_path or os.path.splitext(md_path)[0] + ".docx"
    doc.save(out)
    return out


def main():
    targets = sys.argv[1:]
    if not targets:
        for name in ("RESEARCH_PAPER.md", "THESIS.md"):
            p = os.path.join(HERE, name)
            if os.path.exists(p):
                targets.append(p)
    for t in targets:
        print("wrote:", convert(t))


if __name__ == "__main__":
    main()
